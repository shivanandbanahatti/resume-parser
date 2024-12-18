from docx import Document
from PyPDF2 import PdfReader
import pytesseract
from pdf2image import convert_from_path
import logging
import os
import tempfile

logger = logging.getLogger(__name__)

class DocumentParser:
    def __init__(self):
        # Configure tesseract path
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

    def parse_document(self, file_path: str) -> str:
        """Parse document based on file extension"""
        try:
            if file_path.endswith('.pdf'):
                return self._parse_pdf_with_ocr(file_path)
            elif file_path.endswith('.docx'):
                return self._parse_docx(file_path)
            else:
                raise ValueError("Unsupported file format")
                
        except Exception as e:
            logger.error(f"Error parsing document: {str(e)}")
            raise

    def _parse_pdf_with_ocr(self, file_path: str) -> str:
        """Extract text from PDF using both PDF text extraction and OCR"""
        try:
            text_parts = []
            
            # First try normal PDF text extraction
            reader = PdfReader(file_path)
            for page in reader.pages:
                text_parts.append(page.extract_text())
            
            # Then use OCR for image content
            images = convert_from_path(file_path)
            
            with tempfile.TemporaryDirectory() as temp_dir:
                for i, image in enumerate(images):
                    # Save image temporarily
                    image_path = os.path.join(temp_dir, f'page_{i}.png')
                    image.save(image_path, 'PNG')
                    
                    # Extract text using OCR
                    ocr_text = pytesseract.image_to_string(image_path)
                    text_parts.append(ocr_text)
            
            # Combine all text
            full_text = '\n'.join(text_parts)
            logger.info(f"Extracted text length: {len(full_text)}")
            logger.debug(f"First 200 chars: {full_text[:200]}")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error parsing PDF with OCR: {str(e)}")
            raise

    def _parse_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)
            text = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text.strip())
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(' | '.join(row_text))
            
            return '\n'.join(text)
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise 