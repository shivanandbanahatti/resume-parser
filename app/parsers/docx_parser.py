from docx import Document
import logging

logger = logging.getLogger(__name__)

class DocxParser:
    def extract_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            print("\n=== DEBUG: EXTRACTING TEXT FROM DOCX ===")
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
                    print(f"Paragraph: [{para.text.strip()}]")
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            # Extract text from paragraphs within cells
                            cell_texts = []
                            for paragraph in cell.paragraphs:
                                if paragraph.text.strip():
                                    cell_texts.append(paragraph.text.strip())
                                    print(f"Table cell paragraph: [{paragraph.text.strip()}]")
                            if cell_texts:
                                row_texts.append(' '.join(cell_texts))
                    if row_texts:
                        text_parts.append(' | '.join(row_texts))
            
            # Join with double newlines to preserve structure
            full_text = '\n\n'.join(text_parts)
            
            print("\nExtracted full text:")
            print("---START TEXT---")
            print(full_text)
            print("---END TEXT---\n")
            
            if not full_text.strip():
                raise ValueError("No text content found in the document")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise Exception(f"Error parsing DOCX: {str(e)}")